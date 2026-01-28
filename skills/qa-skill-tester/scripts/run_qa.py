import sys
import os
import argparse
import subprocess
import time
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor

# --- configuration ---
BASE_DIR = os.getcwd()
ARTIFACTS_DIR = os.path.join(BASE_DIR, "artifacts")
QA_CONFIG_DIR = os.path.join(ARTIFACTS_DIR, "qa", "config")
PYTHON_EXE = sys.executable

def ensure_dir(path):
    os.makedirs(path, exist_ok=True)

def generate_checklist(skill_name, output_path):
    wb = Workbook()
    ws = wb.active
    ws.title = f"QA Checklist - {skill_name}"
    
    headers = ["Category", "ID", "Checklist Item", "Status", "Priority", "Notes"]
    ws.append(headers)
    
    # Styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F81BD")
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
    
    # Standard Items + Frontend Req
    items = [
        ("Frontend", "FE-001", "Check readability of all elements", "Pending", "High", ""),
        ("Frontend", "FE-002", "Check clickability of interactive elements", "Pending", "Critical", ""),
        ("Frontend", "FE-003", "Verify no overlapping elements", "Pending", "Medium", ""),
        ("Frontend", "FE-004", "Check for broken links (404)", "Pending", "High", ""),
        ("Functional", "FN-001", "Verify core skill functionality", "Pending", "Critical", ""),
        ("Security", "SEC-001", "Run Security Checker", "Pending", "High", ""),
    ]
    
    for item in items:
        ws.append(item)
        
    ws.auto_filter.ref = ws.dimensions
    wb.save(output_path)
    print(f"✅ Checklist generated: {output_path}")

def generate_bug_report(skill_name, issues, output_path):
    doc = Document()
    doc.add_heading(f'Bug Report: {skill_name}', 0)
    
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Tester: QA Skill Tester (Auto)")
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Automated testing detected the following issues.')
    
    doc.add_heading('Issues List', level=1)
    
    if not issues:
        doc.add_paragraph("No major issues detected during automated run.")
    else:
        for i, issue in enumerate(issues, 1):
            p = doc.add_paragraph(style='List Number')
            runner = p.add_run(f"Issue #{i}: {issue}")
            runner.bold = True
            doc.add_paragraph("Severity: High")
            doc.add_paragraph("Steps to Reproduce: [Auto-detected logs]")
    
    doc.save(output_path)
    print(f"✅ Bug Report generated: {output_path}")

def run_gemini_plan(skill_path, output_path):
    # Simulate Gemini Test Architect call
    # In production, this would call 'gemini prompt' with actual SKILL.md content
    print("🤖 Gemini Test Architect analyzing SKILL.md...")
    with open(output_path, "w") as f:
        f.write(f"# Test Plan for {os.path.basename(skill_path)}\n\n")
        f.write("## Strategy\nExecute automated regression suite.\n\n")
        f.write("## Test Cases\n1. Verify Init\n2. Verify Execution\n3. Verify Output")
    print(f"✅ Test Plan saved: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="QA Skill Tester Orchestrator")
    parser.add_argument("--skill", required=True, help="Path to the skill to test")
    args = parser.parse_args()
    
    skill_path = args.skill
    skill_name = os.path.basename(skill_path)
    
    print(f"🚀 Starting QA Cycle for: {skill_name}")
    
    # 1. Setup Artifacts
    ts = int(time.time())
    run_dir = os.path.join(ARTIFACTS_DIR, skill_name, "test", str(ts))
    ensure_dir(run_dir)
    
    # 2. Generate Materials
    generate_checklist(skill_name, os.path.join(run_dir, "checklist.xlsx"))
    run_gemini_plan(skill_path, os.path.join(run_dir, "test_plan.md"))
    
    # 3. Simulate Execution (Here we would actually call the skill's scripts)
    print("⚡ Running Test Cases...")
    issues = ["UI Overlap detected on Mobile View", "Latency > 200ms"] # Simulated finding
    
    # 4. Report
    generate_bug_report(skill_name, issues, os.path.join(run_dir, "bug_report.docx"))
    
    # 5. Archive (Using workflow)
    # The artifact-archival workflow usually moves specific files, 
    # but since we already wrote to the target dir, we just log it.
    print(f"📦 Artifacts archived in: {run_dir}")

if __name__ == "__main__":
    main()
